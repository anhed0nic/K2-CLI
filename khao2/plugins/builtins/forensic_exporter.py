"""Custom forensic exporter plugin for specialized report formats."""
import csv
import html
import json
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional
from khao2.plugins import (
    ExporterPlugin, PluginMetadata, PluginContext,
    PluginError
)


@dataclass
class ForensicReport:
    """Represents a forensic analysis report."""
    report_id: str
    title: str
    case_number: Optional[str]
    examiner: str
    created_at: float
    evidence_items: List[Dict[str, Any]]
    findings: Dict[str, Any]
    conclusions: List[str]
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class ForensicExporterPlugin(ExporterPlugin):
    """Custom exporter for forensic analysis reports in specialized formats."""

    def __init__(self):
        self.config: Dict[str, Any] = {}
        self.templates: Dict[str, str] = {}

    @property
    def metadata(self) -> PluginMetadata:
        return PluginMetadata(
            name="forensic_exporter",
            version="1.0.0",
            description="Specialized forensic report generation for legal and compliance use",
            author="Khao2 Community",
            plugin_type="exporter",
            entry_point="khao2.plugins.builtins.forensic_exporter.ForensicExporterPlugin",
            config_schema={
                "include_evidence_chain": {
                    "type": "boolean",
                    "default": True,
                    "description": "Include evidence chain of custody information"
                },
                "include_technical_details": {
                    "type": "boolean",
                    "default": True,
                    "description": "Include detailed technical analysis"
                },
                "redact_sensitive_data": {
                    "type": "boolean",
                    "default": False,
                    "description": "Redact sensitive information from reports"
                },
                "default_format": {
                    "type": "string",
                    "enum": ["pdf", "docx", "html", "xml"],
                    "default": "pdf",
                    "description": "Default export format"
                }
            }
        )

    def initialize(self, context: PluginContext) -> None:
        """Initialize the forensic exporter."""
        # Load configuration
        plugin_config = context.config.get('forensic_exporter', {})
        self.config = {
            'include_evidence_chain': plugin_config.get('include_evidence_chain', True),
            'include_technical_details': plugin_config.get('include_technical_details', True),
            'redact_sensitive_data': plugin_config.get('redact_sensitive_data', False),
            'default_format': plugin_config.get('default_format', 'pdf')
        }

        # Load built-in templates
        self._load_templates()

    def cleanup(self) -> None:
        """Clean up resources."""
        pass

    def export(self, data: Any, output_path: Path, **kwargs) -> None:
        """Export forensic report in specified format."""
        format_type = kwargs.get('format', self.config.get('default_format', 'pdf'))

        # Convert data to forensic report format
        if isinstance(data, dict) and 'report_data' in data:
            # Already in report format
            report_data = data['report_data']
        else:
            # Convert from scan results
            report_data = self._convert_to_forensic_report(data, **kwargs)

        if format_type == 'pdf':
            self._export_pdf(report_data, output_path, **kwargs)
        elif format_type == 'docx':
            self._export_docx(report_data, output_path, **kwargs)
        elif format_type == 'html':
            self._export_html(report_data, output_path, **kwargs)
        elif format_type == 'xml':
            self._export_xml(report_data, output_path, **kwargs)
        elif format_type == 'json':
            self._export_json(report_data, output_path, **kwargs)
        else:
            raise PluginError(f"Unsupported export format: {format_type}")

    def _convert_to_forensic_report(self, data: Any, **kwargs) -> ForensicReport:
        """Convert scan results to forensic report format."""
        import uuid

        # Extract case information from kwargs
        case_number = kwargs.get('case_number')
        examiner = kwargs.get('examiner', 'Khao2 Automated Analysis')
        title = kwargs.get('title', f'Forensic Image Analysis Report - {time.strftime("%Y-%m-%d")}')

        evidence_items = []
        findings = {
            "total_scans": 0,
            "suspicious_files": 0,
            "clean_files": 0,
            "anomalies_detected": 0,
            "high_confidence_findings": 0
        }

        conclusions = []

        # Process scan results
        if isinstance(data, list):
            # Batch results
            for item in data:
                if isinstance(item, dict) and 'result' in item:
                    scan_result = item['result']
                    evidence_items.append(self._extract_evidence_item(item))
                    self._update_findings(findings, scan_result)

        elif hasattr(data, 'scans'):
            # Dashboard data
            for scan in data.scans:
                evidence_items.append(self._extract_evidence_item_from_scan(scan))
                self._update_findings(findings, scan)

        # Generate conclusions
        conclusions = self._generate_conclusions(findings)

        return ForensicReport(
            report_id=f"FR_{uuid.uuid4().hex[:8]}",
            title=title,
            case_number=case_number,
            examiner=examiner,
            created_at=time.time(),
            evidence_items=evidence_items,
            findings=findings,
            conclusions=conclusions,
            metadata={
                "generated_by": "Khao2 Forensic Exporter",
                "version": "1.0.0",
                "data_source": "automated_analysis"
            }
        )

    def _extract_evidence_item(self, batch_item: Dict[str, Any]) -> Dict[str, Any]:
        """Extract evidence item from batch result."""
        result = batch_item.get('result', {})

        return {
            "item_id": batch_item.get('scan_id', 'unknown'),
            "file_path": batch_item.get('file_path', 'unknown'),
            "evidence_type": "digital_image",
            "analysis_date": time.time(),
            "findings": {
                "verdict": result.get('verdict', 'unknown'),
                "confidence": result.get('confidence', 0),
                "anomalies": len(result.get('anomalies', []))
            },
            "chain_of_custody": self._generate_chain_of_custody() if self.config['include_evidence_chain'] else None
        }

    def _extract_evidence_item_from_scan(self, scan: Dict[str, Any]) -> Dict[str, Any]:
        """Extract evidence item from scan data."""
        return {
            "item_id": scan.get('scan_id', 'unknown'),
            "file_path": scan.get('file_name', 'unknown'),
            "evidence_type": "digital_image",
            "analysis_date": scan.get('timestamp', time.time()),
            "findings": {
                "verdict": scan.get('verdict', 'unknown'),
                "confidence": scan.get('confidence', 0),
                "anomalies": len(scan.get('anomalies', []))
            },
            "chain_of_custody": self._generate_chain_of_custody() if self.config['include_evidence_chain'] else None
        }

    def _update_findings(self, findings: Dict[str, Any], scan_result: Any) -> None:
        """Update findings summary."""
        findings["total_scans"] += 1

        if hasattr(scan_result, 'static_ai') and scan_result.static_ai:
            ai = scan_result.static_ai
            if ai.possibility_of_steganography > 50:
                findings["suspicious_files"] += 1
            else:
                findings["clean_files"] += 1

            if ai.anomalies:
                findings["anomalies_detected"] += len(ai.anomalies)

            if ai.confidence > 0.8:
                findings["high_confidence_findings"] += 1

    def _generate_conclusions(self, findings: Dict[str, Any]) -> List[str]:
        """Generate conclusions based on findings."""
        conclusions = []

        total = findings.get("total_scans", 0)
        suspicious = findings.get("suspicious_files", 0)

        if total == 0:
            conclusions.append("No evidence items were analyzed.")
            return conclusions

        suspicion_rate = (suspicious / total) * 100

        if suspicion_rate == 0:
            conclusions.append("No suspicious content was detected in the analyzed evidence.")
        elif suspicion_rate < 25:
            conclusions.append(f"Low suspicion rate ({suspicion_rate:.1f}%) detected across evidence items.")
        elif suspicion_rate < 50:
            conclusions.append(f"Moderate suspicion rate ({suspicion_rate:.1f}%) detected across evidence items.")
        else:
            conclusions.append(f"High suspicion rate ({suspicion_rate:.1f}%) detected across evidence items.")

        anomalies = findings.get("anomalies_detected", 0)
        if anomalies > 0:
            conclusions.append(f"{anomalies} anomalies were detected requiring further investigation.")

        return conclusions

    def _generate_chain_of_custody(self) -> Dict[str, Any]:
        """Generate chain of custody information."""
        return {
            "collected_by": "Khao2 Automated Analysis",
            "collection_date": time.strftime("%Y-%m-%d %H:%M:%S"),
            "collection_method": "Digital forensic analysis",
            "preservation_method": "SHA256 hashing and secure storage",
            "analysis_tool": "Khao2 Steganalysis Platform",
            "analysis_date": time.strftime("%Y-%m-%d %H:%M:%S")
        }

    def _export_pdf(self, report: ForensicReport, output_path: Path, **kwargs) -> None:
        """Export report as PDF."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        except ImportError as err:
            raise PluginError("PDF export requires reportlab: pip install reportlab") from err

        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
        )

        story = []

        # Title page
        story.append(Paragraph(report.title, title_style))
        story.append(Spacer(1, 12))

        if report.case_number:
            story.append(Paragraph(f"Case Number: {report.case_number}", styles['Normal']))
            story.append(Spacer(1, 12))

        story.append(Paragraph(f"Examiner: {report.examiner}", styles['Normal']))
        story.append(Paragraph(f"Report Date: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(report.created_at))}", styles['Normal']))
        story.append(Spacer(1, 24))

        # Executive Summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        for conclusion in report.conclusions:
            story.append(Paragraph(f"â€¢ {conclusion}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Findings
        story.append(Paragraph("Key Findings", styles['Heading2']))
        findings_data = [[k.replace('_', ' ').title(), str(v)] for k, v in report.findings.items()]
        findings_table = Table(findings_data)
        findings_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(findings_table)

        doc.build(story)

    def _export_docx(self, report: ForensicReport, output_path: Path, **kwargs) -> None:
        """Export report as DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise PluginError("DOCX export requires python-docx: pip install python-docx")

        doc = Document()
        doc.add_heading(report.title, 0)

        if report.case_number:
            doc.add_paragraph(f"Case Number: {report.case_number}")

        doc.add_paragraph(f"Examiner: {report.examiner}")
        doc.add_paragraph(f"Report Date: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(report.created_at))}")

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        for conclusion in report.conclusions:
            doc.add_paragraph(conclusion, style='List Bullet')

        # Findings
        doc.add_heading('Key Findings', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        # Data rows
        for key, value in report.findings.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)

        doc.save(str(output_path))

    def _export_html(self, report: ForensicReport, output_path: Path, **kwargs) -> None:
        """Export report as HTML."""
        # Escape all user-controlled data to prevent XSS
        esc_title = html.escape(report.title)
        esc_case = html.escape(report.case_number) if report.case_number else ""
        esc_examiner = html.escape(report.examiner)
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>{esc_title}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ background: #f0f0f0; padding: 20px; border-radius: 5px; }}
                .section {{ margin: 20px 0; }}
                .findings {{ background: #e8f4f8; padding: 15px; border-radius: 5px; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{esc_title}</h1>
                {"<p><strong>Case Number:</strong> " + esc_case + "</p>" if esc_case else ""}
                <p><strong>Examiner:</strong> {esc_examiner}</p>
                <p><strong>Report Date:</strong> {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(report.created_at))}</p>
            </div>

            <div class="section">
                <h2>Executive Summary</h2>
                <ul>
                    {"".join(f"<li>{html.escape(conclusion)}</li>" for conclusion in report.conclusions)}
                </ul>
            </div>

            <div class="section findings">
                <h2>Key Findings</h2>
                <table>
                    <tr><th>Metric</th><th>Value</th></tr>
                    {"".join(f"<tr><td>{html.escape(k.replace('_', ' ').title())}</td><td>{html.escape(str(v))}</td></tr>" for k, v in report.findings.items())}
                </table>
            </div>
        </body>
        </html>
        """

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def _export_xml(self, report: ForensicReport, output_path: Path, **kwargs) -> None:
        """Export report as XML."""
        import xml.etree.ElementTree as ET
        from xml.dom import minidom

        root = ET.Element("ForensicReport")
        root.set("id", report.report_id)
        root.set("version", "1.0")

        # Header
        header = ET.SubElement(root, "Header")
        ET.SubElement(header, "Title").text = report.title
        if report.case_number:
            ET.SubElement(header, "CaseNumber").text = report.case_number
        ET.SubElement(header, "Examiner").text = report.examiner
        ET.SubElement(header, "CreatedAt").text = str(report.created_at)

        # Findings
        findings = ET.SubElement(root, "Findings")
        for key, value in report.findings.items():
            finding = ET.SubElement(findings, "Finding")
            finding.set("name", key)
            finding.text = str(value)

        # Conclusions
        conclusions = ET.SubElement(root, "Conclusions")
        for conclusion in report.conclusions:
            ET.SubElement(conclusions, "Conclusion").text = conclusion

        # Pretty print XML
        rough_string = ET.tostring(root, 'utf-8')
        reparsed = minidom.parseString(rough_string)
        pretty_xml = reparsed.toprettyxml(indent="  ")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(pretty_xml)

    def _export_json(self, report: ForensicReport, output_path: Path, **kwargs) -> None:
        """Export report as JSON."""
        report_dict = {
            "report_id": report.report_id,
            "title": report.title,
            "case_number": report.case_number,
            "examiner": report.examiner,
            "created_at": report.created_at,
            "evidence_items": report.evidence_items,
            "findings": report.findings,
            "conclusions": report.conclusions,
            "metadata": report.metadata
        }

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report_dict, f, indent=2, default=str)

    def _load_templates(self) -> None:
        """Load built-in report templates."""
        # Could load from files or define here
        pass


# Plugin metadata for discovery
PLUGIN_METADATA = ForensicExporterPlugin().metadata